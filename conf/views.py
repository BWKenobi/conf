import os
import time
from django.db.models import Q
from datetime import date
import json
from io import BytesIO
from django.core.files import File
from django.core.files.base import ContentFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Pt
from docx.enum.style import WD_STYLE_TYPE

from fpdf import FPDF

from django.conf import settings
from django.templatetags.static import static
from django.contrib.staticfiles.storage import staticfiles_storage
from django.core.files.storage import default_storage

from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.urls import reverse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate, logout, update_session_auth_hash
from django.contrib.auth.models import User

from django.contrib.sites.shortcuts import get_current_site
from django.utils.encoding import force_bytes, force_text
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.template.loader import render_to_string
from django.core.mail import EmailMessage, send_mail
from .tokens import accaunt_activation_token

from profileuser.models import Profile
from coprofile.models import CoProfile
from sections.models import Section

from .forms import UserLoginForm, UserRegistrationForm, OrgRegistrationForm, ChangePasswordForm, CustomPasswordResetForm, CustomSetPasswordForm, SectionForm
from certificates.forms import MakeCertificateForm

class PDF(FPDF):
	pass

def home_view(request):
	members = []
	form = MakeCertificateForm(label_suffix='')
	section_form = SectionForm(label_suffix='')

	users = Profile.objects.all().exclude(username='admin').exclude(user__is_active=False). \
		order_by('surname', 'name', 'name2')#'-moderator_access', '-admin_access', '-speaker', 


	empty_section = False

	for user in users:
		section = ''
		section_pk = ''
		if user.section:
			section = user.section.name
			section_pk = user.section.pk
		elif not user.org_accecc:
			empty_section = True

		member = {
			'pk': user.pk,
			'status_code': user.speaker,
			'name': user.get_full_name(),
			'email': user.user.email,
			'status': user.get_speaker_display(),
			'phone': user.phone,
			'work_place': user.work_place,
			'work_part': user.work_part,
			'position': user.position,
			'degree': user.degree,
			'cert': user.certificate_file,
			'cert_num': user.certificate_num,
			'report_name': user.report_name,
			'report_file': user.report_file,
			'section': section,
			'section_pk': section_pk,
			'org_accecc': user.org_accecc
		}
		members.append(member)
		comembers = CoProfile.objects.filter(lead=user.user)
		if comembers:
			for comember in comembers:
				section = ''
				section_pk = ''
				if user.section:
					section = comember.section.name
					section_pk = user.section.pk
				elif not comember.org_accecc:
					empty_section = True

				member = {
					'name': comember.get_full_name(),
					'email': user.user.email,
					'status': comember.get_speaker_display(),
					'phone': comember.phone,
					'work_place': comember.work_place,
					'work_part': comember.work_part,
					'position': comember.position,
					'degree': comember.degree,
					'cert': comember.certificate_file,
					'cert_num': comember.certificate_num,
					'report_name': comember.report_name,
					'report_file': comember.report_file,
					'section': section,
					'section_pk': section_pk,
					'org_accecc': comember.org_accecc
				}
				members.append(member)
	
	members =  sorted(members, key=lambda i: (i['name']))

	sections = Section.objects.all().order_by('name')
	section_count = {}
	section_name = {}
	super_count = 0
	none_count = 0

	for section in sections:
		section_count[section.pk] = 0
		section_name[section.name] = section.pk

	for member in members:
		if member['org_accecc']:
			super_count += 1
		elif member['section']:
			section_count[section_name[member['section']]] += 1
		else:
			none_count += 1

	count_column = sections.count() + 3
	width_column = 100.0 / count_column

	if request.POST:
		dte = date.today()
		document = Document()
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.PORTRAIT
		section.page_width = Mm(297)
		section.page_height = Mm(210)
		section.left_margin = Mm(30)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)


		document.add_paragraph('Список участников семинара').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		p = document.add_paragraph()
		p.add_run(dte.strftime('%d.%b.%Y')).italic = True
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT


		# p = document.add_paragraph()
		# p.add_run('Орг.комитет').bold = True
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# p.paragraph_format.space_after = 0


		table = document.add_table(rows=1, cols=5)
		table.allow_autifit = False
		table.style = 'TableGrid'
		table.columns[0].width = Mm(10)
		table.columns[1].width = Mm(120)
		table.columns[2].width = Mm(70)
		table.columns[3].width = Mm(30)
		table.columns[4].width = Mm(27)

		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = '№'
		hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[0].width = Mm(10)
		hdr_cells[1].text = 'ФИО, Организация, Должность'
		hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[1].width = Mm(120)
		hdr_cells[2].text = 'Телефон, E-mail'
		hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[2].width = Mm(70)
		hdr_cells[3].text = 'Форма участия'
		hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[3].width = Mm(30)
		hdr_cells[4].text = 'Серт. №'
		hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[4].width = Mm(27)

		count = 1

		for profile in users:
			# if member['section'] == '' and member['org_accecc']:
			row_cells = table.add_row().cells
			row_cells[0].text = str(count)
			row_cells[0].paragraphs[0].runs[0].font.bold = True
			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(10)
			row_cells[1].text = profile.get_full_name()#member['name']
			if profile.degree:#member['degree']:
				row_cells[1].text += ' (' + profile.degree + ')'

			row_cells[1].text += '\n' + profile.work_place#member['work_place'] 

			if profile.work_part:#member['work_part'] :
				row_cells[1].text += ', ' + profile.work_part#member['work_part']

			if profile.position:#member['position']:
				row_cells[1].text += ', ' +  profile.position#member['position']
					
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(120)
			row_cells[2].text = profile.phone + '\n' + profile.user.email#member['phone'] + '\n' + member['email']
			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[2].width = Mm(70)
			row_cells[3].text = profile.get_speaker_display()
			row_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[3].width = Mm(30)
			row_cells[4].text = profile.certificate_num#member['cert_num']
			row_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[4].width = Mm(27)
			count += 1

		# p = document.add_paragraph()
		# p.paragraph_format.space_after = 0

		# sections = Section.objects.all().order_by('name')
		# for section in sections:
		# 	p = document.add_paragraph()
		# 	p.add_run('Секция: ' + section.name).bold = True
		# 	p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	p.paragraph_format.space_after = 0


		# 	table = document.add_table(rows=1, cols=5)
		# 	table.allow_autifit = False
		# 	table.style = 'TableGrid'
		# 	table.columns[0].width = Mm(10)
		# 	table.columns[1].width = Mm(120)
		# 	table.columns[2].width = Mm(70)
		# 	table.columns[3].width = Mm(30)
		# 	table.columns[4].width = Mm(27)

		# 	hdr_cells = table.rows[0].cells
		# 	hdr_cells[0].text = '№'
		# 	hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[0].width = Mm(10)
		# 	hdr_cells[1].text = 'ФИО, Организация, Должность'
		# 	hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[1].width = Mm(120)
		# 	hdr_cells[2].text = 'Телефон, E-mail'
		# 	hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[2].width = Mm(70)
		# 	hdr_cells[3].text = 'Статус'
		# 	hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[3].width = Mm(30)
		# 	hdr_cells[4].text = 'Серт. №'
		# 	hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[4].width = Mm(27)

		# 	count = 1

		# 	for member in members:
		# 		if member['section'] == section.name and not member['org_accecc']:
		# 			row_cells = table.add_row().cells
		# 			row_cells[0].text = str(count)
		# 			row_cells[0].paragraphs[0].runs[0].font.bold = True
		# 			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[0].width = Mm(10)
		# 			row_cells[1].text = member['name']
		# 			if member['degree']:
		# 				row_cells[1].text += ' (' + member['degree'] + ')'

		# 			row_cells[1].text += '\n' + member['work_place'] 

		# 			if member['work_part'] :
		# 				row_cells[1].text += ', ' + member['work_part']

		# 			if member['position']:
		# 				row_cells[1].text += ', ' +  member['position']

		# 			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[1].width = Mm(120)
		# 			row_cells[2].text = member['phone'] + '\n' + member['email']
		# 			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[2].width = Mm(70)
		# 			row_cells[3].text = member['status']
		# 			row_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[3].width = Mm(30)
		# 			row_cells[4].text = member['cert_num']
		# 			row_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[4].width = Mm(27)
		# 			count += 1

		# 	p = document.add_paragraph()
		# 	p.paragraph_format.space_after = 0


		# if empty_section:
		# 	p = document.add_paragraph()
		# 	p.add_run('Без секции').bold = True
		# 	p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	p.paragraph_format.space_after = 0


		# 	table = document.add_table(rows=1, cols=5)
		# 	table.allow_autifit = False
		# 	table.style = 'TableGrid'
		# 	table.columns[0].width = Mm(10)
		# 	table.columns[1].width = Mm(120)
		# 	table.columns[2].width = Mm(70)
		# 	table.columns[3].width = Mm(30)
		# 	table.columns[4].width = Mm(27)

		# 	hdr_cells = table.rows[0].cells
		# 	hdr_cells[0].text = '№'
		# 	hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[0].width = Mm(10)
		# 	hdr_cells[1].text = 'ФИО, Организация, Должность'
		# 	hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[1].width = Mm(120)
		# 	hdr_cells[2].text = 'Телефон, E-mail'
		# 	hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[2].width = Mm(70)
		# 	hdr_cells[3].text = 'Статус'
		# 	hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[3].width = Mm(30)
		# 	hdr_cells[4].text = 'Серт. №'
		# 	hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		# 	hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 	hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 	hdr_cells[4].width = Mm(27)

		# 	count = 1

		# 	for member in members:
		# 		if member['section'] == '' and not member['org_accecc']:
		# 			row_cells = table.add_row().cells
		# 			row_cells[0].text = str(count)
		# 			row_cells[0].paragraphs[0].runs[0].font.bold = True
		# 			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[0].width = Mm(10)
		# 			row_cells[1].text = member['name']
		# 			if member['degree']:
		# 				row_cells[1].text += ' (' + member['degree'] + ')'

		# 			row_cells[1].text += '\n' + member['work_place'] 

		# 			if member['work_part'] :
		# 				row_cells[1].text += ', ' + member['work_part']

		# 			if member['position']:
		# 				row_cells[1].text += ', ' +  member['position']
							
		# 			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[1].width = Mm(120)
		# 			row_cells[2].text = member['phone'] + '\n' + member['email']
		# 			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[2].width = Mm(70)
		# 			row_cells[3].text = member['status']
		# 			row_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[3].width = Mm(30)
		# 			row_cells[4].text = member['cert_num']
		# 			row_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# 			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		# 			row_cells[4].width = Mm(27)
		# 			count += 1




		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=List (' + dte.strftime('%d-%b-%Y') + ').docx'
		document.save(response)

		return response




	args = {
		'users': users,
		'members': members,
		'form': form,
		'section_form': section_form,
		'sections': sections,
		'section_count': section_count,
		'super_count': super_count,
		'none_count': none_count,

		'count_column': count_column,
		'width_column': width_column
	}
	return render(request, 'index.html', args)


def policy_view(request):
	return render(request, 'policy.html')


def login_view(request):
	form = UserLoginForm(request.POST or None)
	next_ = request.GET.get('next')
	modal = False

	if form.is_valid():
		username = request.POST.get('email').lower()
		password = request.POST.get('password')
		user = authenticate(username = username.strip(), password = password.strip())
		
		login(request, user)
		next_post = request.POST.get('next')
		redirect_path = next_ or next_post or '/'


		return redirect(redirect_path)

	args = {
		'form': form
	}
	return render(request, 'login.html', args)


def logout_view(request):
	logout(request)
	return redirect('home')


def register_view(request):
	if request.method=='POST':
		user_form = UserRegistrationForm(request.POST)
		if user_form.is_valid():
			new_user = user_form.save(commit=False)
			new_user.username = new_user.email
			new_user.is_active = False
			new_user.set_password(user_form.cleaned_data['password'])
			new_user.save()

			new_user.profile.name = user_form.cleaned_data['name']
			new_user.profile.name2 = user_form.cleaned_data['name2']
			new_user.profile.surname = user_form.cleaned_data['surname']
			new_user.profile.work_place = user_form.cleaned_data['work_place']

			new_user.profile.phone = user_form.cleaned_data['phone']
			new_user.profile.work_part = user_form.cleaned_data['work_part']
			new_user.profile.position = user_form.cleaned_data['position']
			new_user.profile.degree = user_form.cleaned_data['degree']
			# new_user.profile.speaker = user_form.cleaned_data['speaker']
			
			# section = Section.objects.get(pk = user_form.cleaned_data['section'])
			# new_user.profile.section = section
			new_user.profile.save()


			current_site = get_current_site(request)
			protocol = 'http'
			if request.is_secure():
				protocol = 'https'

			mail_subject = 'Активация аккаунта'
			to_email = new_user.email
			#if '127.0.0.1' in current_site.domain:
			uid = urlsafe_base64_encode(force_bytes(new_user.pk))
			#else:
			#	uid = urlsafe_base64_encode(force_bytes(new_user.pk)).decode()

			token = accaunt_activation_token.make_token(new_user)

			message = render_to_string('acc_active_email.html', {'user': new_user, 'domain': current_site.domain,\
				'uid': uid, \
				'token': token,\
				'protocol': protocol,\
				'email': to_email})

			message_html = render_to_string('acc_active_email_html.html', {'user': new_user, 'domain': current_site.domain,\
				'uid': uid, \
				'token': token,\
				'protocol': protocol,\
				'email': to_email})

			t = send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

			return render(request, 'confirm_email.html')

		return render(request, 'register.html', {'form': user_form})

	user_form = UserRegistrationForm()
	return render(request, 'register.html', {'form': user_form})


def superviser_view(request):
	if request.method=='POST':
		user_form = OrgRegistrationForm(request.POST)
		if user_form.is_valid():
			new_user = user_form.save(commit=False)
			new_user.username = new_user.email
			new_user.is_active = False
			new_user.set_password(user_form.cleaned_data['password'])
			new_user.save()

			new_user.profile.name = user_form.cleaned_data['name']
			new_user.profile.name2 = user_form.cleaned_data['name2']
			new_user.profile.surname = user_form.cleaned_data['surname']
			new_user.profile.work_place = user_form.cleaned_data['work_place']

			new_user.profile.phone = user_form.cleaned_data['phone']
			new_user.profile.work_part = user_form.cleaned_data['work_part']
			new_user.profile.position = user_form.cleaned_data['position']
			new_user.profile.degree = user_form.cleaned_data['degree']
			new_user.profile.speaker = user_form.cleaned_data['speaker']
			
			new_user.profile.org_accecc = True
			new_user.profile.save()


			current_site = get_current_site(request)
			protocol = 'http'
			if request.is_secure():
				protocol = 'https'

			mail_subject = 'Активация аккаунта'
			to_email = new_user.email
			#if '127.0.0.1' in current_site.domain:
			uid = urlsafe_base64_encode(force_bytes(new_user.pk))
			#else:
			#	uid = urlsafe_base64_encode(force_bytes(new_user.pk)).decode()

			token = accaunt_activation_token.make_token(new_user)

			message = render_to_string('acc_active_email.html', {'user': new_user, 'domain': current_site.domain,\
				'uid': uid, \
				'token': token,\
				'protocol': protocol,\
				'email': to_email})

			message_html = render_to_string('acc_active_email_html.html', {'user': new_user, 'domain': current_site.domain,\
				'uid': uid, \
				'token': token,\
				'protocol': protocol,\
				'email': to_email})

			t = send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

			return render(request, 'confirm_email.html')

		return render(request, 'superviser.html', {'form': user_form})

	user_form = OrgRegistrationForm()
	return render(request, 'superviser.html', {'form': user_form})


def activate(request, uidb64, token):
	try:
		uid = force_text(urlsafe_base64_decode(uidb64))
		user = User.objects.get(pk=uid)
	except(TypeError, ValueError, OverflowError, User.DoesNotExist):
		user = None

	if user and accaunt_activation_token.check_token(user, token):
		user.is_active = True
		user.save()

		user_count = User.objects.filter(profile__section = user.profile.section, is_active = True).count()
		if user.profile.section:
			if user_count > user.profile.section.count:
				user.profile.section = None
				user.profile.save()

		login(request, user)
		return redirect('home')

	return render(request, 'failed_email.html')


def change_password(request):
	username = request.user.username

	if request.method == 'POST':
		form = ChangePasswordForm(request.POST, username=username)

		if form.is_valid():
			user = request.user
			user.set_password(form.cleaned_data['newpassword1'])
			user.save()
			update_session_auth_hash(request, user)

			return redirect('profiles:view_edit_profile')

		args = {
			'form': form
		}
		return render(request, 'change_password.html', args)

	form = ChangePasswordForm(username=username)

	args = {
		'form': form,
		'menu': 'profile',
	}
	return render(request, 'change_password.html', args)


# --------------------------------
#           Для ajax'а
# --------------------------------
def change_admin_access(request):
	user_id = request.GET.get('user_id')
	access = request.GET.get('access')
	if access=='true':
		access = True
	else:
		access = False
	
	profile = Profile.objects.get(id = user_id)
	profile.admin_access = access
	profile.save()
	
	return HttpResponse(json.dumps('Статус изменен.'))


def change_moderate_access(request):
	user_id = request.GET.get('user_id')
	access = request.GET.get('access')
	if access=='true':
		access = True
	else:
		access = False
	
	profile = Profile.objects.get(id = user_id)
	profile.moderator_access = access
	profile.save()
	
	return HttpResponse(json.dumps('Статус изменен.'))


def send_info_message(request):
	users = User.objects.all().exclude(username='admin').exclude(is_active=False)
	count = 0

	for user in users:
		mail_subject = 'Информационное письмо'
		to_email = user.email
		sex = user.profile.sex()
		sex_valid = user.profile.sex_valid()
		
		message = render_to_string('info_email.html', {'sex_valid': sex_valid, 'sex': sex, 'name': user.profile.get_io_name()})

		message_html = render_to_string('info_email_html.html', {'sex_valid': sex_valid, 'sex': sex, 'name': user.profile.get_io_name()})

		email = EmailMessage(mail_subject, message, settings.EMAIL_HOST_USER, [to_email])

		docfile = default_storage.open(user.profile.certificate_file.name, 'r')

		email.attach_file(docfile.name)

		email.send()

		#send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

		count += 1

		if count==5:
			count = 0
			time.sleep(1.5)

		cousers = CoProfile.objects.filter(lead=user)
		if cousers:
			for couser in cousers:
				mail_subject = 'Информационное письмо'
				to_email = user.email
				sex = couser.sex()
				sex_valid = couser.sex_valid()
				
				message = render_to_string('info_email.html', {'sex_valid': sex_valid, 'sex': sex, 'name': couser.get_io_name()})

				message_html = render_to_string('info_email_html.html', {'sex_valid': sex_valid, 'sex': sex, 'name': couser.get_io_name()})

				email = EmailMessage(mail_subject, message, settings.EMAIL_HOST_USER, [to_email])

				docfile = default_storage.open(couser.certificate_file.name, 'r')

				email.attach_file(docfile.name)

				email.send()

				#send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

				count += 1

				if count==5:
					count = 0
					time.sleep(1.5)

	return HttpResponse(json.dumps('Сообщения разосланы.'))